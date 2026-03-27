from __future__ import annotations

import re
import shutil
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / "output" / "doc"


NAVY = RGBColor(0x10, 0x2A, 0x43)
TEAL = RGBColor(0x1F, 0x6F, 0x78)
SAND = RGBColor(0xD6, 0xB3, 0x7A)
GRAY = RGBColor(0x5B, 0x65, 0x6F)


@dataclass
class DocJob:
    source: Path
    target: Path
    title_for_footer: str
    cover: bool = True
    page_break_before_primary_heading: bool = False


DOCS = [
    DocJob(
        source=ROOT / "docs" / "setters" / "manual-unificado.md",
        target=OUTPUT_DIR / "manual-setters.docx",
        title_for_footer="Manual de setters",
        cover=True,
        page_break_before_primary_heading=False,
    ),
    DocJob(
        source=ROOT / "docs" / "setters" / "hoja-rapida-operativa.md",
        target=OUTPUT_DIR / "hoja-rapida-setters.docx",
        title_for_footer="Hoja rapida para setters",
        cover=False,
        page_break_before_primary_heading=False,
    ),
    DocJob(
        source=ROOT / "docs" / "setters" / "presentacion-capacitacion.md",
        target=OUTPUT_DIR / "presentacion-setters-guion.docx",
        title_for_footer="Presentacion para setters",
        cover=True,
        page_break_before_primary_heading=True,
    ),
]


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_page_margins(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)


def set_footer(section, text: str) -> None:
    footer = section.footer
    footer.is_linked_to_previous = False
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.font.name = "Aptos"
    run.font.size = Pt(8)
    run.font.color.rgb = GRAY


def configure_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor(0x23, 0x2F, 0x34)

    title = document.styles["Title"]
    title.font.name = "Georgia"
    title.font.size = Pt(28)
    title.font.bold = True
    title.font.color.rgb = NAVY

    heading_1 = document.styles["Heading 1"]
    heading_1.font.name = "Georgia"
    heading_1.font.size = Pt(18)
    heading_1.font.bold = True
    heading_1.font.color.rgb = NAVY

    heading_2 = document.styles["Heading 2"]
    heading_2.font.name = "Aptos"
    heading_2.font.size = Pt(13)
    heading_2.font.bold = True
    heading_2.font.color.rgb = TEAL

    heading_3 = document.styles["Heading 3"]
    heading_3.font.name = "Aptos"
    heading_3.font.size = Pt(11)
    heading_3.font.bold = True
    heading_3.font.color.rgb = NAVY

    for style_name in ["List Bullet", "List Number"]:
        style = document.styles[style_name]
        style.font.name = "Aptos"
        style.font.size = Pt(10.5)


def add_cover(document: Document, title: str, subtitle: str | None) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(72)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.style = document.styles["Title"]
    paragraph.add_run(title)

    if subtitle:
        sub = document.add_paragraph()
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub.paragraph_format.space_before = Pt(10)
        sub.paragraph_format.space_after = Pt(28)
        run = sub.add_run(subtitle)
        run.font.name = "Aptos"
        run.font.size = Pt(12)
        run.italic = True
        run.font.color.rgb = GRAY

    band = document.add_table(rows=1, cols=1)
    band.alignment = WD_ALIGN_PARAGRAPH.CENTER
    band.autofit = True
    cell = band.cell(0, 0)
    set_cell_shading(cell, "F7F4EE")
    cell.text = "Nivora | Kit comercial para setters"
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cell.paragraphs[0].runs[0]
    run.font.name = "Aptos"
    run.font.size = Pt(10)
    run.font.bold = True
    run.font.color.rgb = TEAL

    document.add_page_break()


def add_title_at_top(document: Document, title: str, subtitle: str | None) -> None:
    paragraph = document.add_paragraph(style="Title")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.add_run(title)

    if subtitle:
        sub = document.add_paragraph()
        sub.paragraph_format.space_after = Pt(14)
        run = sub.add_run(subtitle)
        run.font.name = "Aptos"
        run.font.size = Pt(10)
        run.italic = True
        run.font.color.rgb = GRAY


def add_callout(document: Document, text: str) -> None:
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_ALIGN_PARAGRAPH.LEFT
    table.autofit = True
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F7F4EE")
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.font.name = "Aptos"
    run.font.size = Pt(10)
    run.italic = True
    run.font.color.rgb = NAVY
    document.add_paragraph()


def add_paragraph(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.15
    paragraph.add_run(text)


def add_list_item(document: Document, text: str, numbered: bool) -> None:
    style = "List Number" if numbered else "List Bullet"
    paragraph = document.add_paragraph(style=style)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.add_run(text)


def split_front_matter(lines: list[str]) -> tuple[str, str | None, list[str]]:
    title = "Documento"
    subtitle = None
    cursor = 0

    if lines and lines[0].startswith("# "):
        title = lines[0][2:].strip()
        cursor = 1

    while cursor < len(lines) and not lines[cursor].strip():
        cursor += 1

    if cursor < len(lines) and lines[cursor].startswith("> "):
        subtitle = lines[cursor][2:].strip()
        cursor += 1

    return title, subtitle, lines[cursor:]


def render_markdown(document: Document, lines: list[str], page_break_before_primary_heading: bool) -> None:
    paragraph_buffer: list[str] = []
    seen_primary_heading = False

    def flush_paragraph() -> None:
        if paragraph_buffer:
            add_paragraph(document, " ".join(part.strip() for part in paragraph_buffer))
            paragraph_buffer.clear()

    for raw_line in lines:
        line = raw_line.rstrip()
        stripped = line.strip()

        if not stripped:
            flush_paragraph()
            continue

        if stripped == "---":
            flush_paragraph()
            document.add_page_break()
            continue

        if stripped.startswith("> "):
            flush_paragraph()
            add_callout(document, stripped[2:].strip())
            continue

        if stripped.startswith("## "):
            flush_paragraph()
            if page_break_before_primary_heading and seen_primary_heading:
                document.add_page_break()
            seen_primary_heading = True
            heading = document.add_heading(stripped[3:].strip(), level=1)
            heading.paragraph_format.space_before = Pt(4)
            heading.paragraph_format.space_after = Pt(8)
            continue

        if stripped.startswith("### "):
            flush_paragraph()
            heading = document.add_heading(stripped[4:].strip(), level=2)
            heading.paragraph_format.space_before = Pt(4)
            heading.paragraph_format.space_after = Pt(4)
            continue

        if stripped.startswith("- "):
            flush_paragraph()
            add_list_item(document, stripped[2:].strip(), numbered=False)
            continue

        if re.match(r"^\d+\.\s+", stripped):
            flush_paragraph()
            content = re.sub(r"^\d+\.\s+", "", stripped)
            add_list_item(document, content, numbered=True)
            continue

        paragraph_buffer.append(stripped)

    flush_paragraph()


def generate_doc(job: DocJob) -> None:
    markdown = job.source.read_text(encoding="utf-8").splitlines()
    title, subtitle, body_lines = split_front_matter(markdown)

    document = Document()
    set_page_margins(document)
    configure_styles(document)
    set_footer(document.sections[0], job.title_for_footer)
    document.core_properties.title = title

    if job.cover:
        add_cover(document, title, subtitle)
        new_section = document.sections[-1]
        new_section.start_type = WD_SECTION_START.NEW_PAGE
        set_footer(new_section, job.title_for_footer)
    else:
        add_title_at_top(document, title, subtitle)

    render_markdown(document, body_lines, job.page_break_before_primary_heading)
    job.target.parent.mkdir(parents=True, exist_ok=True)
    document.save(job.target)


def write_manifest() -> None:
    manifest = OUTPUT_DIR / "README.md"
    manifest.write_text(
        "\n".join(
            [
                "# Entregables para setters",
                "",
                "Archivos generados:",
                "",
                "- manual-setters.docx",
                "- hoja-rapida-setters.docx",
                "- presentacion-setters-guion.docx",
                "",
                "Fuentes maestras:",
                "",
                "- docs/setters/manual-unificado.md",
                "- docs/setters/hoja-rapida-operativa.md",
                "- docs/setters/presentacion-capacitacion.md",
                "",
                "Nota: el entorno actual no tiene herramientas instaladas para exportar o previsualizar PDF automaticamente.",
            ]
        ),
        encoding="utf-8",
    )


def copy_sources() -> None:
    for job in DOCS:
        target_copy = OUTPUT_DIR / job.source.name
        shutil.copy2(job.source, target_copy)


def main() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    for job in DOCS:
        generate_doc(job)
    copy_sources()
    write_manifest()


if __name__ == "__main__":
    main()
